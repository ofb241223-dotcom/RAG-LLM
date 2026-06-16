from __future__ import annotations

import html
import re
import time
import zipfile
import subprocess
import tempfile
from io import BytesIO
from pathlib import Path
from shutil import which
from uuid import uuid4

import requests
from charset_normalizer import from_bytes
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from pypdf import PdfReader
import xlrd

from rag_service.documents.formats import DOC, DOCX, PDF, TXT, XLS, XLSX, normalize_format
from rag_service.documents.parser import DocumentParser, ParsedDocument, ParserRegistry
from rag_service.observability import RequestLogger


class DocumentParsingError(ValueError):
    """Raised when bytes have the right extension but cannot be parsed."""


class MinerUDocumentParser:
    def __init__(
        self,
        *,
        api_token: str | None,
        base_url: str = "https://mineru.net/api/v4",
        model_version: str = "vlm",
        timeout_seconds: int = 300,
        poll_interval_seconds: float = 2.0,
        session: requests.Session | None = None,
        request_logger: RequestLogger | None = None,
    ) -> None:
        self.api_token = api_token
        self.base_url = base_url.rstrip("/")
        self.model_version = model_version
        self.timeout_seconds = timeout_seconds
        self.poll_interval_seconds = poll_interval_seconds
        self.session = session or requests.Session()
        self.request_logger = request_logger

    def available(self) -> bool:
        return bool(self.api_token and self.api_token.strip())

    def parse_document(self, content: bytes, *, source_name: str) -> str:
        if not self.available():
            raise DocumentParsingError("MINERU_API_TOKEN is not configured.")

        try:
            batch_id = self._upload_file(content, source_name=source_name)
            zip_url = self._wait_for_zip(batch_id)
            return self._download_markdown(zip_url)
        except requests.RequestException as error:
            raise DocumentParsingError(f"MinerU request failed: {error}") from error
        except (KeyError, TypeError, ValueError, zipfile.BadZipFile) as error:
            raise DocumentParsingError(f"MinerU response could not be parsed: {error}") from error

    def parse_pdf(self, content: bytes, *, source_name: str) -> str:
        return self.parse_document(content, source_name=source_name)

    def _headers(self) -> dict[str, str]:
        return {"Authorization": f"Bearer {self.api_token}"}

    def _upload_file(self, content: bytes, *, source_name: str) -> str:
        safe_name = Path(source_name).name or f"document-{uuid4().hex}.pdf"
        payload = {
            "files": [{"name": safe_name, "data_id": f"rag-{uuid4().hex}"}],
            "model_version": self.model_version,
            "enable_table": True,
            "enable_formula": True,
            "language": "ch",
        }
        response = self.session.post(
            f"{self.base_url}/file-urls/batch",
            headers={**self._headers(), "Content-Type": "application/json"},
            json=payload,
            timeout=30,
        )
        self._raise_for_mineru(response)
        self._log("POST", "/file-urls/batch", _status_code(response), source_name)
        body = response.json()
        data = body["data"]
        upload_url = data["file_urls"][0]
        put_response = self.session.put(upload_url, data=content, timeout=min(self.timeout_seconds, 120))
        put_response.raise_for_status()
        self._log("PUT", "upload-url", _status_code(put_response), safe_name)
        return str(data["batch_id"])

    def _wait_for_zip(self, batch_id: str) -> str:
        deadline = time.monotonic() + self.timeout_seconds
        last_state = ""
        last_error = ""
        while time.monotonic() < deadline:
            response = self.session.get(
                f"{self.base_url}/extract-results/batch/{batch_id}",
                headers=self._headers(),
                timeout=30,
            )
            self._raise_for_mineru(response)
            self._log("GET", "/extract-results/batch", _status_code(response), batch_id)
            body = response.json()
            data = body.get("data") or {}
            results = data.get("extract_result") or data.get("extract_results") or []
            if results:
                result = results[0]
                last_state = str(result.get("state") or data.get("state") or "")
                last_error = str(result.get("err_msg") or result.get("error_msg") or data.get("err_msg") or "")
                zip_url = result.get("full_zip_url") or result.get("zip_url")
                if last_state == "done" and zip_url:
                    return str(zip_url)
                if last_state == "failed":
                    raise DocumentParsingError(f"MinerU parsing failed: {last_error or batch_id}")
            else:
                last_state = str(data.get("state") or "")
            time.sleep(self.poll_interval_seconds)
        raise DocumentParsingError(f"MinerU parsing timed out: {last_state or 'unknown'} {last_error}".strip())

    def _download_markdown(self, zip_url: str) -> str:
        response = self.session.get(zip_url, timeout=60)
        response.raise_for_status()
        self._log("GET", "result-zip", _status_code(response), "download markdown")
        with zipfile.ZipFile(BytesIO(response.content)) as archive:
            markdown_name = next((name for name in archive.namelist() if name.endswith("full.md")), None)
            if markdown_name is None:
                markdown_name = next((name for name in archive.namelist() if name.endswith(".md")), None)
            if markdown_name is None:
                raise DocumentParsingError("MinerU result zip does not contain Markdown output.")
            return _clean_mineru_markdown(archive.read(markdown_name).decode("utf-8", errors="replace"))

    @staticmethod
    def _raise_for_mineru(response: requests.Response) -> None:
        response.raise_for_status()
        body = response.json()
        code = body.get("code")
        if code not in (0, "0", None):
            raise DocumentParsingError(f"MinerU API returned code {code}: {body.get('msg') or body.get('message')}")

    def _log(self, method: str, path: str, status: int, summary: str) -> None:
        if self.request_logger is not None:
            self.request_logger({
                "direction": "PROVIDER",
                "service": "MinerU",
                "method": method,
                "path": path,
                "status": status,
                "summary": summary,
            })


def _clean_mineru_markdown(text: str) -> str:
    cleaned = html.unescape(text)
    cleaned = re.sub(r"!\[[^\]]*]\([^)]*\)", "", cleaned)
    cleaned = re.sub(r"</td>\s*<td>", " | ", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"<tr>\s*<td>", "\n", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"</td>\s*</tr>", "\n", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"</?(?:table|tbody|thead|details|summary)[^>]*>", "\n", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"<br\s*/?>", "\n", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"<[^>]+>", "", cleaned)
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _status_code(response: object) -> int:
    value = getattr(response, "status_code", 200)
    return value if isinstance(value, int) else 200


def _parse_with_mineru(
    mineru_parser: MinerUDocumentParser | None,
    content: bytes,
    *,
    source_name: str,
) -> str | None:
    if mineru_parser is None or not mineru_parser.available():
        return None
    try:
        text = mineru_parser.parse_document(content, source_name=source_name)
    except DocumentParsingError:
        return None
    return text.strip() or None


class TxtDocumentParser:
    supported_formats = (TXT,)

    def can_parse(self, format_name: str) -> bool:
        return normalize_format(format_name) in self.supported_formats

    def parse(self, content: bytes, *, source_name: str, format_name: str) -> ParsedDocument:
        matches = from_bytes(content)
        match = next((candidate for candidate in matches if candidate.language == "Chinese"), None)
        if match is None:
            match = matches.best()
        text = str(match) if match is not None else content.decode("utf-8", errors="replace")
        return ParsedDocument(text=text.strip(), format=format_name, source_name=source_name)


class PdfDocumentParser:
    supported_formats = (PDF,)

    def __init__(self, *, ocr_timeout_seconds: int = 180, ocr_dpi: int = 180, mineru_parser: MinerUDocumentParser | None = None) -> None:
        self.ocr_timeout_seconds = ocr_timeout_seconds
        self.ocr_dpi = ocr_dpi
        self.mineru_parser = mineru_parser

    def can_parse(self, format_name: str) -> bool:
        return normalize_format(format_name) in self.supported_formats

    def parse(self, content: bytes, *, source_name: str, format_name: str) -> ParsedDocument:
        mineru_text = _parse_with_mineru(self.mineru_parser, content, source_name=source_name)
        if mineru_text is not None:
            metadata = {}
            page_count = self._try_count_pages(content)
            if page_count is not None:
                metadata["page_count"] = page_count
            return ParsedDocument(
                text=mineru_text,
                format=format_name,
                source_name=source_name,
                metadata=metadata,
            )

        try:
            reader = PdfReader(BytesIO(content))
            page_text = [(page.extract_text() or "").strip() for page in reader.pages]
        except Exception as error:  # pypdf raises several parser-specific exceptions.
            raise DocumentParsingError(f"Failed to parse PDF: {source_name}") from error

        text = "\n\n".join(part for part in page_text if part)
        if not text.strip():
            text = self._parse_scanned_pdf(content, source_name=source_name)

        return ParsedDocument(
            text=text.strip(),
            format=format_name,
            source_name=source_name,
            metadata={"page_count": len(reader.pages)},
        )

    def _try_count_pages(self, content: bytes) -> int | None:
        try:
            return len(PdfReader(BytesIO(content)).pages)
        except Exception:
            return None

    def _parse_scanned_pdf(self, content: bytes, *, source_name: str) -> str:
        text = _parse_with_mineru(self.mineru_parser, content, source_name=source_name)
        if text is not None:
            return text
        return self._ocr_pdf(content, source_name=source_name)

    def _ocr_pdf(self, content: bytes, *, source_name: str) -> str:
        if which("pdftoppm") is None or which("tesseract") is None:
            return ""

        with tempfile.TemporaryDirectory(prefix="rag-pdf-ocr-") as tmpdir:
            tmpdir_path = Path(tmpdir)
            input_path = tmpdir_path / "document.pdf"
            output_prefix = tmpdir_path / "page"
            input_path.write_bytes(content)

            convert = subprocess.run(
                [
                    "pdftoppm",
                    "-r",
                    str(self.ocr_dpi),
                    "-png",
                    str(input_path),
                    str(output_prefix),
                ],
                capture_output=True,
                text=True,
                timeout=self.ocr_timeout_seconds,
                check=False,
            )
            if convert.returncode != 0:
                detail = convert.stderr.strip() or convert.stdout.strip()
                raise DocumentParsingError(f"PDF has no text layer and OCR conversion failed: {detail or source_name}")

            page_images = sorted(tmpdir_path.glob("page-*.png"))
            if not page_images:
                raise DocumentParsingError(f"PDF has no text layer and OCR conversion produced no pages: {source_name}")

            page_text: list[str] = []
            per_page_timeout = max(15, self.ocr_timeout_seconds // max(1, len(page_images)))
            for image_path in page_images:
                ocr = subprocess.run(
                    [
                        "tesseract",
                        str(image_path),
                        "stdout",
                        "-l",
                        "chi_sim+eng",
                        "--psm",
                        "6",
                    ],
                    capture_output=True,
                    text=True,
                    timeout=per_page_timeout,
                    check=False,
                )
                if ocr.returncode != 0:
                    detail = ocr.stderr.strip() or ocr.stdout.strip()
                    raise DocumentParsingError(f"PDF has no text layer and OCR failed: {detail or source_name}")
                if ocr.stdout.strip():
                    page_text.append(ocr.stdout.strip())

            return "\n\n".join(page_text)


class DocxDocumentParser:
    supported_formats = (DOCX,)

    def __init__(self, *, mineru_parser: MinerUDocumentParser | None = None) -> None:
        self.mineru_parser = mineru_parser

    def can_parse(self, format_name: str) -> bool:
        return normalize_format(format_name) in self.supported_formats

    def parse(self, content: bytes, *, source_name: str, format_name: str) -> ParsedDocument:
        mineru_text = _parse_with_mineru(self.mineru_parser, content, source_name=source_name)
        if mineru_text is not None:
            return ParsedDocument(text=mineru_text, format=format_name, source_name=source_name)

        try:
            document = Document(BytesIO(content))
        except Exception as error:
            raise DocumentParsingError(f"Failed to parse DOCX: {source_name}") from error

        return ParsedDocument(text=_extract_docx_text(document), format=format_name, source_name=source_name)


def _extract_docx_text(document: DocxDocument) -> str:
    lines: list[str] = []
    for block in _iter_docx_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                lines.append(text)
        else:
            for row in block.rows:
                cells = [_normalize_docx_cell(cell.text) for cell in row.cells]
                line = " | ".join(cell for cell in cells if cell)
                if line:
                    lines.append(line)
    return "\n".join(lines)


def _iter_docx_blocks(document: DocxDocument):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def _normalize_docx_cell(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


class ExcelDocumentParser:
    supported_formats = (XLSX, XLS)

    def __init__(self, *, mineru_parser: MinerUDocumentParser | None = None) -> None:
        self.mineru_parser = mineru_parser

    def can_parse(self, format_name: str) -> bool:
        return normalize_format(format_name) in self.supported_formats

    def parse(self, content: bytes, *, source_name: str, format_name: str) -> ParsedDocument:
        mineru_text = _parse_with_mineru(self.mineru_parser, content, source_name=source_name)
        if mineru_text is not None:
            return ParsedDocument(text=mineru_text, format=format_name, source_name=source_name)

        normalized = normalize_format(format_name)
        if normalized == XLSX:
            text = self._parse_xlsx(content, source_name=source_name)
        elif normalized == XLS:
            text = self._parse_xls(content, source_name=source_name)
        else:
            raise DocumentParsingError(f"Unsupported Excel format: {format_name}")
        return ParsedDocument(text=text, format=format_name, source_name=source_name)

    def _parse_xlsx(self, content: bytes, *, source_name: str) -> str:
        try:
            workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
        except Exception as error:
            raise DocumentParsingError(f"Failed to parse XLSX: {source_name}") from error

        sections: list[str] = []
        for worksheet in workbook.worksheets:
            rows = [
                " | ".join(_normalize_excel_cell(value) for value in row if _normalize_excel_cell(value))
                for row in worksheet.iter_rows(values_only=True)
            ]
            rows = [row for row in rows if row]
            if rows:
                sections.append("\n".join([f"Sheet: {worksheet.title}", *rows]))
        return "\n\n".join(sections)

    def _parse_xls(self, content: bytes, *, source_name: str) -> str:
        try:
            workbook = xlrd.open_workbook(file_contents=content)
        except Exception as error:
            raise DocumentParsingError(f"Failed to parse XLS: {source_name}") from error

        sections: list[str] = []
        for sheet in workbook.sheets():
            rows: list[str] = []
            for row_index in range(sheet.nrows):
                cells = [_normalize_excel_cell(sheet.cell_value(row_index, col_index)) for col_index in range(sheet.ncols)]
                line = " | ".join(cell for cell in cells if cell)
                if line:
                    rows.append(line)
            if rows:
                sections.append("\n".join([f"Sheet: {sheet.name}", *rows]))
        return "\n\n".join(sections)


def _normalize_excel_cell(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return re.sub(r"\s+", " ", str(value)).strip()


class DocDocumentParser:
    supported_formats = (DOC,)

    def __init__(self, *, timeout_seconds: int = 30, mineru_parser: MinerUDocumentParser | None = None) -> None:
        self.timeout_seconds = timeout_seconds
        self.mineru_parser = mineru_parser

    def can_parse(self, format_name: str) -> bool:
        return normalize_format(format_name) in self.supported_formats

    def parse(self, content: bytes, *, source_name: str, format_name: str) -> ParsedDocument:
        mineru_text = _parse_with_mineru(self.mineru_parser, content, source_name=source_name)
        if mineru_text is not None:
            return ParsedDocument(text=mineru_text, format=format_name, source_name=source_name)

        source_path = Path(source_name)
        safe_name = source_path.name or "document.doc"
        if not safe_name.lower().endswith(".doc"):
            safe_name = f"{source_path.stem or 'document'}.doc"

        with tempfile.TemporaryDirectory(prefix="rag-doc-") as tmpdir:
            tmpdir_path = Path(tmpdir)
            input_path = tmpdir_path / safe_name
            input_path.write_bytes(content)

            args = [
                "libreoffice",
                "--headless",
                "--convert-to",
                "txt:Text (encoded):UTF8",
                "--outdir",
                str(tmpdir_path),
                str(input_path),
            ]
            try:
                result = subprocess.run(
                    args,
                    capture_output=True,
                    text=True,
                    timeout=self.timeout_seconds,
                    check=False,
                )
            except FileNotFoundError as error:
                raise DocumentParsingError("LibreOffice is required to parse .doc files.") from error
            except subprocess.TimeoutExpired as error:
                raise DocumentParsingError(f"LibreOffice timed out while parsing DOC: {source_name}") from error

            if result.returncode != 0:
                detail = result.stderr.strip() or result.stdout.strip()
                raise DocumentParsingError(f"LibreOffice failed to parse DOC: {detail}")

            output_path = input_path.with_suffix(".txt")
            if not output_path.exists():
                candidates = list(tmpdir_path.glob("*.txt"))
                if not candidates:
                    raise DocumentParsingError(f"LibreOffice did not produce text for DOC: {source_name}")
                output_path = candidates[0]

            return ParsedDocument(
                text=output_path.read_text(encoding="utf-8", errors="replace").strip(),
                format=format_name,
                source_name=source_name,
            )


def create_default_parser_registry(
    *,
    mineru_api_token: str | None = None,
    mineru_api_base_url: str = "https://mineru.net/api/v4",
    mineru_enabled: bool = False,
    mineru_model_version: str = "vlm",
    mineru_timeout_seconds: int = 300,
    mineru_poll_interval_seconds: float = 2.0,
    request_logger: RequestLogger | None = None,
) -> ParserRegistry:
    mineru_parser = None
    if mineru_enabled:
        mineru_parser = MinerUDocumentParser(
            api_token=mineru_api_token,
            base_url=mineru_api_base_url,
            model_version=mineru_model_version,
            timeout_seconds=mineru_timeout_seconds,
            poll_interval_seconds=mineru_poll_interval_seconds,
            request_logger=request_logger,
        )
    return ParserRegistry(
        (
            TxtDocumentParser(),
            PdfDocumentParser(mineru_parser=mineru_parser),
            DocxDocumentParser(mineru_parser=mineru_parser),
            DocDocumentParser(mineru_parser=mineru_parser),
            ExcelDocumentParser(mineru_parser=mineru_parser),
        )
    )


__all__ = [
    "DocDocumentParser",
    "DocxDocumentParser",
    "DocumentParsingError",
    "ExcelDocumentParser",
    "MinerUDocumentParser",
    "PdfDocumentParser",
    "TxtDocumentParser",
    "create_default_parser_registry",
]
