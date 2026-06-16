import subprocess
import zipfile
from io import BytesIO
from pathlib import Path

import pytest

from rag_service.documents.parsers import (
    DocDocumentParser,
    DocxDocumentParser,
    ExcelDocumentParser,
    MinerUDocumentParser,
    PdfDocumentParser,
    TxtDocumentParser,
    _clean_mineru_markdown,
    create_default_parser_registry,
)


def test_txt_parser_detects_non_utf8_encoding() -> None:
    content = "自然语言处理综述".encode("gb18030")

    parsed = TxtDocumentParser().parse(content, source_name="notes.txt", format_name="txt")

    assert parsed.text == "自然语言处理综述"
    assert parsed.format == "txt"
    assert parsed.source_name == "notes.txt"


def test_docx_parser_extracts_paragraphs() -> None:
    from docx import Document

    buffer = BytesIO()
    document = Document()
    document.add_paragraph("第一段")
    document.add_paragraph("第二段")
    document.save(buffer)

    parsed = DocxDocumentParser().parse(buffer.getvalue(), source_name="自然语言处理综述.docx", format_name="docx")

    assert parsed.text == "第一段\n第二段"
    assert parsed.format == "docx"


def test_docx_parser_extracts_tables_when_falling_back_locally() -> None:
    from docx import Document

    buffer = BytesIO()
    document = Document()
    document.add_paragraph("名单")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "姓名"
    table.cell(0, 1).text = "状态"
    table.cell(1, 0).text = "张三"
    table.cell(1, 1).text = "推荐"
    document.save(buffer)

    parsed = DocxDocumentParser().parse(buffer.getvalue(), source_name="名单.docx", format_name="docx")

    assert parsed.text == "名单\n姓名 | 状态\n张三 | 推荐"


def test_docx_parser_prefers_mineru_when_available() -> None:
    class FakeMinerU:
        def available(self) -> bool:
            return True

        def parse_document(self, content: bytes, *, source_name: str) -> str:
            assert content == b"docx-bytes"
            assert source_name == "report.docx"
            return "MinerU Word Markdown"

    parsed = DocxDocumentParser(mineru_parser=FakeMinerU()).parse(
        b"docx-bytes",
        source_name="report.docx",
        format_name="docx",
    )

    assert parsed.text == "MinerU Word Markdown"


def test_pdf_parser_uses_pypdf_reader(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakePage:
        def __init__(self, text: str) -> None:
            self._text = text

        def extract_text(self) -> str:
            return self._text

    class FakeReader:
        def __init__(self, _: BytesIO) -> None:
            self.pages = [FakePage("第一页"), FakePage("第二页")]

    monkeypatch.setattr("rag_service.documents.parsers.PdfReader", FakeReader)

    parsed = PdfDocumentParser().parse(b"%PDF", source_name="chapter.pdf", format_name="pdf")

    assert parsed.text == "第一页\n\n第二页"
    assert parsed.metadata["page_count"] == 2


def test_pdf_parser_prefers_mineru_even_when_pdf_has_text_layer(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakePage:
        def extract_text(self) -> str:
            return "本地文本层"

    class FakeReader:
        def __init__(self, _: BytesIO) -> None:
            self.pages = [FakePage()]

    class FakeMinerU:
        def available(self) -> bool:
            return True

        def parse_document(self, content: bytes, *, source_name: str) -> str:
            assert content == b"%PDF"
            assert source_name == "chapter.pdf"
            return "MinerU 结构化 Markdown"

    monkeypatch.setattr("rag_service.documents.parsers.PdfReader", FakeReader)

    parsed = PdfDocumentParser(mineru_parser=FakeMinerU()).parse(b"%PDF", source_name="chapter.pdf", format_name="pdf")

    assert parsed.text == "MinerU 结构化 Markdown"
    assert parsed.metadata["page_count"] == 1


def test_pdf_parser_uses_ocr_when_pdf_has_no_text_layer(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakePage:
        def extract_text(self) -> str:
            return ""

    class FakeReader:
        def __init__(self, _: BytesIO) -> None:
            self.pages = [FakePage()]

    def fake_run(args: list[str], **_: object) -> subprocess.CompletedProcess[str]:
        if args[0] == "pdftoppm":
            output_prefix = Path(args[-1])
            (output_prefix.parent / "page-1.png").write_bytes(b"png")
            return subprocess.CompletedProcess(args=args, returncode=0, stdout="", stderr="")
        if args[0] == "tesseract":
            return subprocess.CompletedProcess(args=args, returncode=0, stdout="扫描版 PDF 文字", stderr="")
        raise AssertionError(f"Unexpected command: {args}")

    monkeypatch.setattr("rag_service.documents.parsers.PdfReader", FakeReader)
    monkeypatch.setattr("rag_service.documents.parsers.which", lambda command: f"/usr/bin/{command}")
    monkeypatch.setattr("rag_service.documents.parsers.subprocess.run", fake_run)

    parsed = PdfDocumentParser().parse(b"%PDF", source_name="scan.pdf", format_name="pdf")

    assert parsed.text == "扫描版 PDF 文字"
    assert parsed.metadata["page_count"] == 1


def test_pdf_parser_uses_mineru_before_local_ocr_when_pdf_has_no_text_layer(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakePage:
        def extract_text(self) -> str:
            return ""

    class FakeReader:
        def __init__(self, _: BytesIO) -> None:
            self.pages = [FakePage()]

    class FakeMinerU:
        def available(self) -> bool:
            return True

        def parse_document(self, content: bytes, *, source_name: str) -> str:
            return self.parse_pdf(content, source_name=source_name)

        def parse_pdf(self, content: bytes, *, source_name: str) -> str:
            assert content == b"%PDF"
            assert source_name == "scan.pdf"
            return "MinerU Markdown 文本"

    monkeypatch.setattr("rag_service.documents.parsers.PdfReader", FakeReader)
    monkeypatch.setattr("rag_service.documents.parsers.which", lambda _command: None)

    parsed = PdfDocumentParser(mineru_parser=FakeMinerU()).parse(b"%PDF", source_name="scan.pdf", format_name="pdf")

    assert parsed.text == "MinerU Markdown 文本"


def test_mineru_pdf_parser_uploads_polls_and_reads_full_markdown() -> None:
    class FakeResponse:
        def __init__(self, *, body: dict[str, object] | None = None, content: bytes = b"") -> None:
            self._body = body or {}
            self.content = content

        def json(self) -> dict[str, object]:
            return self._body

        def raise_for_status(self) -> None:
            return None

    class FakeSession:
        def __init__(self) -> None:
            self.put_content = b""
            self.polls = 0

        def post(self, url: str, **kwargs: object) -> FakeResponse:
            assert url == "https://mineru.test/file-urls/batch"
            assert kwargs["json"] == {
                "files": [{"name": "scan.pdf", "data_id": kwargs["json"]["files"][0]["data_id"]}],  # type: ignore[index]
                "model_version": "vlm",
                "enable_table": True,
                "enable_formula": True,
                "language": "ch",
            }
            return FakeResponse(body={"code": 0, "data": {"batch_id": "batch-1", "file_urls": ["https://upload.test/scan.pdf"]}})

        def put(self, url: str, **kwargs: object) -> FakeResponse:
            assert url == "https://upload.test/scan.pdf"
            self.put_content = kwargs["data"]  # type: ignore[assignment]
            return FakeResponse()

        def get(self, url: str, **_: object) -> FakeResponse:
            if url.endswith("/extract-results/batch/batch-1"):
                self.polls += 1
                return FakeResponse(
                    body={
                        "code": 0,
                        "data": {
                            "extract_result": [
                                {
                                    "state": "done",
                                    "full_zip_url": "https://download.test/result.zip",
                                }
                            ]
                        },
                    }
                )
            assert url == "https://download.test/result.zip"
            buffer = BytesIO()
            with zipfile.ZipFile(buffer, "w") as archive:
                archive.writestr("full.md", "# 标题\n\n| 姓名 | 状态 |\n| --- | --- |\n| 张三 | 推荐 |")
            return FakeResponse(content=buffer.getvalue())

    session = FakeSession()
    parser = MinerUDocumentParser(
        api_token="token",
        base_url="https://mineru.test",
        timeout_seconds=10,
        poll_interval_seconds=0,
        session=session,
    )

    text = parser.parse_pdf(b"%PDF", source_name="scan.pdf")

    assert session.put_content == b"%PDF"
    assert session.polls == 1
    assert "| 张三 | 推荐 |" in text


def test_clean_mineru_markdown_turns_html_tables_into_readable_rows() -> None:
    text = "<table><tr><td>姓名</td><td>状态</td></tr><tr><td>张三</td><td>推荐</td></tr></table>"

    cleaned = _clean_mineru_markdown(text)

    assert "姓名 | 状态" in cleaned
    assert "张三 | 推荐" in cleaned
    assert "<td>" not in cleaned


def test_doc_parser_converts_with_libreoffice_using_argument_array(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    captured_args: list[str] = []

    def fake_run(args: list[str], **_: object) -> subprocess.CompletedProcess[str]:
        captured_args.extend(args)
        output_dir = Path(args[args.index("--outdir") + 1])
        (output_dir / "legacy.txt").write_text("旧版 Word 文本", encoding="utf-8")
        return subprocess.CompletedProcess(args=args, returncode=0, stdout="", stderr="")

    monkeypatch.setattr("rag_service.documents.parsers.subprocess.run", fake_run)

    parsed = DocDocumentParser(timeout_seconds=1).parse(b"doc-bytes", source_name="legacy.doc", format_name="doc")

    assert parsed.text == "旧版 Word 文本"
    assert captured_args[:5] == ["libreoffice", "--headless", "--convert-to", "txt:Text (encoded):UTF8", "--outdir"]


def test_doc_parser_prefers_mineru_when_available() -> None:
    class FakeMinerU:
        def available(self) -> bool:
            return True

        def parse_document(self, content: bytes, *, source_name: str) -> str:
            assert content == b"doc-bytes"
            assert source_name == "legacy.doc"
            return "MinerU DOC Markdown"

    parsed = DocDocumentParser(mineru_parser=FakeMinerU(), timeout_seconds=1).parse(
        b"doc-bytes",
        source_name="legacy.doc",
        format_name="doc",
    )

    assert parsed.text == "MinerU DOC Markdown"


def test_xlsx_parser_extracts_sheets_when_falling_back_locally() -> None:
    from openpyxl import Workbook

    buffer = BytesIO()
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "推荐名单"
    worksheet.append(["姓名", "专业", "状态"])
    worksheet.append(["张三", "软件工程", "推荐"])
    workbook.save(buffer)

    parsed = ExcelDocumentParser().parse(buffer.getvalue(), source_name="名单.xlsx", format_name="xlsx")

    assert parsed.text == "Sheet: 推荐名单\n姓名 | 专业 | 状态\n张三 | 软件工程 | 推荐"
    assert parsed.format == "xlsx"


def test_excel_parser_prefers_mineru_when_available() -> None:
    class FakeMinerU:
        def available(self) -> bool:
            return True

        def parse_document(self, content: bytes, *, source_name: str) -> str:
            assert content == b"xlsx-bytes"
            assert source_name == "table.xlsx"
            return "MinerU Excel Markdown"

    parsed = ExcelDocumentParser(mineru_parser=FakeMinerU()).parse(
        b"xlsx-bytes",
        source_name="table.xlsx",
        format_name="xlsx",
    )

    assert parsed.text == "MinerU Excel Markdown"


def test_default_parser_registry_handles_all_supported_formats() -> None:
    registry = create_default_parser_registry()

    for format_name in ("pdf", "txt", "docx", "doc", "xlsx", "xls"):
        assert registry.get_parser(format_name).can_parse(format_name)
